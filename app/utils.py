from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# fname ="Путь к файлу"
#
# # Генерируете новый ключ (или берете ранее сгенерированный)
# key = RSA.generate(1024, os.urandom)
# # Получаете хэш файла
# h = SHA256.new()
# with open(fname, "rb") as f:
#     for chunk in iter(lambda: f.read(4096), b""):
#         h.update(chunk)
#
# # Подписываете хэш
# signature = pkcs1_15.new(key).sign(h)
#
# # Получаете открытый ключ из закрытого
# pubkey = key.publickey()
#
# # Пересылаете пользователю файл, публичный ключ и подпись
# # На стороне пользователя заново вычисляете хэш файла (опущено) и сверяете подпись
# pkcs1_15.new(pubkey).verify(h, signature)
#
# # Отличающийся хэш не должен проходить проверку
# pkcs1_15.new(pubkey).verify(SHA256.new(b'test'), signature) # raise ValueError("Invalid signature")


def generate_report(data):
    document = Document()

    head = document.add_heading(
        'Квитанция на оплату закупок компанией "{}" с {} по {}'.format(data['Компания'], data['Дата начала отсчета'],
                                                                       data['Дата конца отсчета']), 0)
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = document.add_paragraph('Уважаемый(ая) ')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
    p.add_run(data['Фамилия'] + ' ' + data['Имя'] + ' ' + data['Отчество'] + ' ').bold = True
    p.add_run('просим вас оплатить закупки, произведенные компанией ')
    p.add_run('"' + data['Компания'] + '"' + ', ').bold = True
    p.add_run('в которой для нас вы являеетесь контактным лицом в период с ')
    p.add_run(data['Дата начала отсчета'] + ' ').bold = True
    p.add_run('по ')
    p.add_run(data['Дата конца отсчета'] + '.').bold = True

    document.add_paragraph()
    document.add_paragraph('Прилагаю список товаров:')

    records = data['Информация о закупках']

    table = document.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Дата'
    hdr_cells[1].text = 'Продукт'
    hdr_cells[2].text = 'Количество'
    hdr_cells[3].text = 'Цена'
    hdr_cells[4].text = 'Итого'

    for deal in records.keys():
        row_cells = table.add_row().cells
        row_cells[0].text = deal
        row_cells[1].text = records[deal][0]['Продукт']
        row_cells[2].text = records[deal][0]['Количество'].__str__()
        row_cells[3].text = records[deal][0]['Цена продажи'].__str__()
        row_cells[4].text = records[deal][0]['Сумма'].__str__()
        for index, element in enumerate(records[deal]):
            if index > 0:
                row_cells = table.add_row().cells
                row_cells[1].text = records[deal][index]['Продукт']
                row_cells[2].text = records[deal][index]['Количество'].__str__()
                row_cells[3].text = records[deal][index]['Цена продажи'].__str__()
                row_cells[4].text = records[deal][index]['Сумма'].__str__()

    document.add_paragraph()
    p = document.add_paragraph('Итого общая сумма вашей задолженности составляет ')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
    p.add_run(data['Общая сумма закупок'].__str__() + ' рублей').bold = True

    document.add_page_break()
    document.save('demo.docx')
